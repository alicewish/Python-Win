import time

from docx import Document

# from docx.shared import Inches

start_time = time.time()  # 初始时间戳
document = Document()
file_path = '/Users/alicewish/我的坚果云/中文测试.docx'

document.add_heading('文档标题', 0)

p = document.add_paragraph('一个普通的段落,有')
p.add_run('粗体').bold = True
p.add_run('和')
p.add_run('斜体。').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
# for item in recordset:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(item.qty)
#     row_cells[1].text = str(item.id)
#     row_cells[2].text = item.desc

document.add_page_break()

document.save(file_path)
# ================运行时间计时================
run_time = time.time() - start_time
if run_time < 60:  # 两位小数的秒
    print("耗时:{:.2f}秒".format(run_time))
elif run_time < 3600:  # 分秒取整
    print("耗时:{:.0f}分{:.0f}秒".format(run_time // 60, run_time % 60))
else:  # 时分秒取整
    print("耗时:{:.0f}时{:.0f}分{:.0f}秒".format(run_time // 3600, run_time % 3600 // 60, run_time % 60))
